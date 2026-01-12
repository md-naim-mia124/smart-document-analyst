# # app.py


# import os
# import re
# import io
# import zipfile
# import unicodedata
# import streamlit as st
# from typing import List, Dict, Optional

# from config.settings import UPLOAD_DIR, SUPPORTED_EXTENSIONS
# from modules.doc_loader import extract_pages
# from modules.toc_extractor import extract_toc
# from modules.chatbot import Chatbot
# from modules.translator import Translator

# # -----------------------------
# # Helpers
# # -----------------------------
# def find_section_for_chunk(chunk_text: str, toc: List[Dict], full_text: str) -> str:
#     chunk_pos = full_text.find(chunk_text)
#     if chunk_pos == -1:
#         return "Unknown"
#     current_section = "Introduction"
#     for section in toc:
#         title = section.get("title", "")
#         try:
#             pos = full_text.lower().find(title.lower())
#             if pos != -1 and pos <= chunk_pos:
#                 current_section = title
#             else:
#                 break
#         except Exception:
#             continue
#     return current_section

# def _safe_build_index(doc_id: str, file_path: str, toc: List[Dict]):
#     from modules.vectorstore import VectorStore
#     from modules.embedder import Embedder
#     store, embedder = VectorStore(doc_id), Embedder()
#     pages = extract_pages(file_path)
#     full_text = "\n\n".join(p_text for _, p_text in pages)
#     chunk_dicts, texts_to_embed = [], []
#     for page_num, page_text in pages:
#         for chunk in store.split(page_text):
#             section_name = find_section_for_chunk(chunk, toc, full_text)
#             chunk_dicts.append({"text": chunk, "page": page_num, "section": section_name})
#             texts_to_embed.append(chunk)
#     if not texts_to_embed:
#         raise ValueError("Document is empty or could not be chunked.")
#     vectors = embedder.encode(texts_to_embed)
#     store.build_index(vectors, chunk_dicts)
#     return len(chunk_dicts)

# @st.cache_resource
# def get_chatbot(doc_id, file_path):
#     return Chatbot(doc_id, file_path)

# @st.cache_resource
# def get_translator():
#     model_path = os.getenv("TRANSLATION_MODEL", "models/m2m100_418M")
#     return Translator(model_path)

# def render_toc(toc, level=0):
#     for entry in toc:
#         indent = "&nbsp;" * 4 * (entry.get("level", 0))
#         st.sidebar.markdown(
#             f"{indent}- **{entry['title']}**" + (f" (p.{entry.get('page')})" if entry.get("page") else ""),
#             unsafe_allow_html=True,
#         )
#         if entry.get("children"):
#             render_toc(entry["children"], level + 1)

# # --- Output cleaning / sanitizing ---
# _DROP_PATTERNS = [
#     r"\bThe provided evidence does not contain\b.*",
#     r"\bThe provided evidence does not\b.*",
# ]

# def clean_answer(raw: str) -> str:
#     """Clean LLM answer while preserving markdown tables and structure."""
#     if not isinstance(raw, str) or not raw.strip():
#         return raw
#     txt = raw.replace("[/ ", "").replace("[/]", "").replace("[ /]", "")
#     out_lines = []
#     for ln in txt.splitlines():
#         # Preserve markdown table rows
#         if ln.lstrip().startswith("|"):
#             out_lines.append(ln)
#             continue
#         # Drop boilerplate evidence disclaimers
#         if any(re.search(p, ln, re.I) for p in _DROP_PATTERNS):
#             continue
#         # Remove bold wrappers at line starts only
#         ln = re.sub(r"^\s*\*\*(.*?)\s*\*\*\s*", r"\1", ln)
#         out_lines.append(ln)
#     txt = "\n".join(out_lines)
#     txt = re.sub(r"\n{3,}", "\n\n", txt).strip()
#     return txt

# _CLEAN_TAGS = [
#     r"\[E\d+\]",   # evidence tags like [E1]
#     r"\[/\s*",     # stray [/ markers
# ]

# def strip_artifacts(text: str) -> str:
#     """Remove evidence tags and stray markers; keep tables intact."""
#     if not isinstance(text, str):
#         return text
#     out = text
#     for p in _CLEAN_TAGS:
#         out = re.sub(p, "", out)
#     out = re.sub(r"\s{2,}", " ", out)
#     out = re.sub(r"\n{3,}", "\n\n", out)
#     return out.strip()

# def _slugify(text: str, maxlen: int = 80) -> str:
#     if not text:
#         return ""
#     value = unicodedata.normalize("NFKD", text).encode("ascii", "ignore").decode("ascii")
#     value = re.sub(r"[^\w\s-]+", "", value).strip().lower()
#     value = re.sub(r"[-\s]+", "_", value)
#     return value[:maxlen] if maxlen else value

# def _export_suffix(item: Dict) -> str:
#     """Map history item to filename suffix."""
#     typ = (item.get("type") or "").lower()
#     lang = (item.get("lang") or "en").lower()

#     if typ.startswith("q&a"):
#         base = "qa"
#     elif typ.startswith("summarize") or "summary" in typ:
#         base = "summary"
#     elif typ.startswith("novelty"):
#         base = "novelty"
#     elif typ.startswith("review"):
#         base = "review"
#     else:
#         base = _slugify(item.get("type") or "result")

#     if "translated" in typ:
#         m = re.search(r"translated\s*—\s*([A-Z]+)", item.get("type", ""), re.I)
#         if m:
#             lang_name = m.group(1).lower()
#         else:
#             lang_name = lang
#         return f"{base}_{lang_name}_translated"

#     return base

# def _item_to_text(item: Dict) -> str:
#     """Serialize an item for export (txt/docx/pdf)."""
#     parts = []
#     # For Q&A items, content already contains Question/Answer; don't prepend "Q:"
#     if item.get("type", "").lower().startswith("q&a"):
#         parts.append(item.get("content", ""))
#     else:
#         if item.get("question"):
#             parts.append(f"Q: {item['question']}")
#         parts.append(item.get("content", ""))
#     if "latency" in item:
#         parts.append(f"\n---\nGenerated in {item['latency']:.2f} seconds.")
#     return "\n\n".join([p for p in parts if p])

# def _build_txt_bytes(item: Dict) -> bytes:
#     return _item_to_text(item).encode("utf-8")

# def _build_docx_bytes(item: Dict) -> bytes:
#     from docx import Document
#     from docx.shared import Pt
#     doc = Document()
#     if item.get("type"):
#         p = doc.add_paragraph()
#         run = p.add_run(item["type"])
#         run.bold = True
#         run.font.size = Pt(14)
#     # For non-Q&A items, keep a top "Q:" line if present
#     if item.get("question") and not item.get("type", "").lower().startswith("q&a"):
#         doc.add_paragraph(f"Q: {item['question']}")
#     for line in _item_to_text(item).splitlines():
#         doc.add_paragraph(line)
#     buf = io.BytesIO()
#     doc.save(buf)
#     buf.seek(0)
#     return buf.read()

# def _build_pdf_bytes(item: Dict) -> Optional[bytes]:
#     try:
#         from reportlab.lib.pagesizes import A4
#         from reportlab.pdfgen import canvas
#         from reportlab.lib.units import cm
#     except Exception:
#         return None
#     buf = io.BytesIO()
#     c = canvas.Canvas(buf, pagesize=A4)
#     width, height = A4
#     x, y = 2 * cm, height - 2 * cm

#     title = item.get("type", "Result")
#     c.setFont("Helvetica-Bold", 12)
#     c.drawString(x, y, title)
#     y -= 0.8 * cm

#     # For non-Q&A items, include a "Q:" header
#     if item.get("question") and not item.get("type", "").lower().startswith("q&a"):
#         c.setFont("Helvetica-Bold", 10)
#         c.drawString(x, y, f"Q: {item['question']}")
#         y -= 0.7 * cm

#     c.setFont("Helvetica", 10)
#     textobj = c.beginText(x, y)
#     for line in _item_to_text(item).splitlines():
#         textobj.textLine(line if line else " ")
#         if textobj.getY() < 2 * cm:
#             c.drawText(textobj)
#             c.showPage()
#             textobj = c.beginText(x, height - 2 * cm)
#             c.setFont("Helvetica", 10)
#     c.drawText(textobj)
#     c.showPage()
#     c.save()
#     buf.seek(0)
#     return buf.read()

# def _build_bytes_for_format(item: Dict, fmt: str) -> Optional[bytes]:
#     fmt = fmt.lower()
#     if fmt == "txt":
#         return _build_txt_bytes(item)
#     elif fmt == "docx":
#         return _build_docx_bytes(item)
#     elif fmt == "pdf":
#         return _build_pdf_bytes(item)
#     else:
#         return _build_txt_bytes(item)

# def format_qa_block(question: str, answer: str) -> str:
#     q = (question or "").strip()
#     a = (answer or "").strip()

#     # 1) Strip common Q/A headers the model might add
#     a = re.sub(r"^\s*(Q|Question)\s*:\s*.*\n+", "", a, flags=re.IGNORECASE)
#     a = re.sub(r"^\s*(A|Answer)\s*:\s*", "", a, flags=re.IGNORECASE)

#     # 2) Try to remove the literal question if it’s echoed at the start
#     #    Allow optional **bold**, quotes, trailing punctuation, colon/dash separators.
#     escaped_q = re.escape(q.rstrip("?.!\"'`*:_- \t"))
#     leading_q_patterns = [
#         rf"^\s*\*\*{escaped_q}\*\*\s*[:\-–—]?\s*",
#         rf'^\s*"{escaped_q}"\s*[:\-–—]?\s*',
#         rf"^\s*{escaped_q}\s*[?.!:]*\s*[:\-–—]?\s*",
#     ]
#     for pat in leading_q_patterns:
#         new_a = re.sub(pat, "", a, flags=re.IGNORECASE)
#         if new_a != a:
#             a = new_a
#             break

#     # 3) As a fallback: if the very first sentence == question (ignoring case/punct), drop it
#     def _norm(s: str) -> str:
#         s = s.lower().strip()
#         s = re.sub(r"[^\w\s]", " ", s)   # remove punctuation for compare
#         s = re.sub(r"\s+", " ", s)
#         return s

#     # first sentence of answer
#     first_sentence_match = re.split(r"(?<=[.!?])\s+", a, maxsplit=1)
#     first_sent = first_sentence_match[0] if first_sentence_match else a
#     if _norm(first_sent) == _norm(q) or _norm(first_sent.rstrip("?.!:")) == _norm(q):
#         a = a[len(first_sent):].lstrip()

#     # 4) Collapse extra blank lines
#     a = re.sub(r"\n{3,}", "\n\n", a).strip()

#     return f"**Question**\n\n{q}\n\n**Answer**\n\n{a}\n"

# # -----------------------------
# # Page & State
# # -----------------------------
# st.set_page_config(page_title="Smart Document Analysis", layout="wide")
# st.title("📄 Smart Document Analyst")

# st.session_state.setdefault("docs", [])
# st.session_state.setdefault("history", {})
# st.session_state.setdefault("active_doc_id", None)
# st.session_state.setdefault("export_mode", False)
# st.session_state.setdefault("active_doc_base", None)

# translator = get_translator()

# # -----------------------------
# # Sidebar Controls
# # -----------------------------
# st.sidebar.subheader("⚙️ Generation Controls")
# st.session_state.setdefault("max_new_tokens", 900)
# st.session_state.setdefault("target_words", 600)
# st.session_state.setdefault("verbosity", "detailed")
# st.session_state.setdefault("strict_grounding", True)
# st.session_state.setdefault("show_evidence", True)

# st.session_state["max_new_tokens"] = st.sidebar.number_input(
#     "Max new tokens", min_value=128, max_value=4096, step=64, value=st.session_state["max_new_tokens"]
# )
# st.session_state["target_words"] = st.sidebar.slider(
#     "Target length (words)", min_value=150, max_value=1500, step=50, value=st.session_state["target_words"]
# )
# st.session_state["verbosity"] = st.sidebar.selectbox(
#     "Writing style", ["concise", "balanced", "detailed", "exhaustive"],
#     index=["concise", "balanced", "detailed", "exhaustive"].index(st.session_state["verbosity"]),
# )
# st.session_state["strict_grounding"] = st.sidebar.checkbox(
#     "Strict grounding (evidence-only, no invention)", value=st.session_state["strict_grounding"]
# )
# st.session_state["show_evidence"] = st.sidebar.checkbox("Show evidence panel", value=st.session_state["show_evidence"])

# # -----------------------------
# # Uploader & Indexing
# # -----------------------------
# uploaded = st.file_uploader("Upload your document", type=[e.lstrip(".") for e in SUPPORTED_EXTENSIONS])
# if uploaded:
#     doc_id, _ext = os.path.splitext(uploaded.name)
#     st.session_state["active_doc_base"] = doc_id
#     if not any(d["doc_id"] == doc_id for d in st.session_state.docs):
#         path = os.path.join(UPLOAD_DIR, uploaded.name)
#         os.makedirs(UPLOAD_DIR, exist_ok=True)
#         with open(path, "wb") as f:
#             f.write(uploaded.getbuffer())
#         with st.spinner(f"Processing and indexing {uploaded.name}..."):
#             toc = extract_toc(path)
#             st.session_state.docs.append({"file_path": path, "doc_id": doc_id, "toc": toc})
#             st.session_state.history[doc_id] = []
#             try:
#                 _safe_build_index(doc_id, path, toc)
#                 st.sidebar.success(f"✅ Indexed {uploaded.name}.")
#                 st.session_state.active_doc_id = doc_id
#             except Exception as e:
#                 st.sidebar.error(f"Index build failed: {e}")
#     else:
#         if st.session_state.active_doc_id != doc_id:
#             st.session_state.active_doc_id = doc_id

# # -----------------------------
# # Main App
# # -----------------------------
# if not st.session_state.active_doc_id:
#     st.info("Please upload a document to begin.")
#     st.stop()

# doc_id = st.session_state.active_doc_id
# try:
#     current_doc = next(d for d in st.session_state.docs if d["doc_id"] == doc_id)
# except StopIteration:
#     st.error("Active document not found. Please re-upload.", icon="🚨")
#     st.stop()

# if not st.session_state.get("active_doc_base"):
#     st.session_state["active_doc_base"] = current_doc["doc_id"]

# file_path = current_doc["file_path"]
# doc_history = st.session_state.history.setdefault(doc_id, [])
# doc_base = st.session_state["active_doc_base"]

# # -----------------------------
# # Sidebar: Table of Contents (ONLY here)
# # -----------------------------
# if st.sidebar.checkbox("📚 Show Table of Contents", value=True, key="show_toc"):
#     toc = current_doc.get("toc", [])
#     if toc:
#         st.sidebar.subheader("📚 Table of Contents")
#         render_toc(toc)
#     else:
#         st.sidebar.info("No Table of Contents extracted for this document.")

# # -----------------------------
# # Generate
# # -----------------------------
# st.subheader("💬 Chat & Generate")
# mode = st.radio("Select Mode", ["General Q&A", "Summarize", "Novelty", "Review"], horizontal=True, key=f"mode_{doc_id}")
# bot = get_chatbot(doc_id, file_path)

# # --- Generation Logic (Q&A only) ---
# progress_placeholder = st.empty()

# if mode == "General Q&A":
#     q = st.text_input("Your question:", key=f"q_{doc_id}")
#     if st.button("Ask", key=f"ask_{doc_id}") and q:
#         with st.spinner("🤖 Thinking..."):
#             # FIX: Pass the placeholder to the ask function
#             ans, elapsed = bot.ask(
#                 q, 
#                 max_new_tokens=st.session_state.max_new_tokens, 
#                 target_words=st.session_state.target_words, 
#                 verbosity=st.session_state.verbosity, 
#                 strict=st.session_state.strict_grounding,
#                 progress_placeholder=progress_placeholder
#             )
#         # store per-doc history (kept consistent with the rest of the app)
#         doc_history.append({"type": "Q&A", "question": q, "content": ans, "latency": elapsed})

# # --- Generation Logic (Summarize only) ---
# progress_placeholder = st.empty()

# if mode == "Summarize":
#     prompts = {"Summarize": "summarize", "Novelty": "novelty", "Review": "review"}
#     if st.button(f"Run {mode}", key=f"run_{mode}_{doc_id}"):
#         ans, elapsed = bot.ask(
#             prompts[mode], 
#             max_new_tokens=st.session_state.max_new_tokens, 
#             target_words=st.session_state.target_words, 
#             verbosity=st.session_state.verbosity, 
#             strict=st.session_state.strict_grounding,
#             progress_placeholder=progress_placeholder
#         )
#         # store per-doc history
#         doc_history.append({"type": mode, "content": ans, "latency": elapsed})

# # --- Generation Logic (Novelty only) ---
# progress_placeholder = st.empty()

# if mode == "Novelty":
#     prompts = {"Summarize": "summarize", "Novelty": "novelty", "Review": "review"}
#     if st.button(f"Run {mode}", key=f"run_{mode}_{doc_id}"):
#         ans, elapsed = bot.ask(
#             prompts[mode], 
#             max_new_tokens=st.session_state.max_new_tokens, 
#             target_words=st.session_state.target_words, 
#             verbosity=st.session_state.verbosity, 
#             strict=st.session_state.strict_grounding,
#             progress_placeholder=progress_placeholder
#         )
#         # store per-doc history
#         doc_history.append({"type": mode, "content": ans, "latency": elapsed})

# # --- Generation Logic (Review only) ---
# progress_placeholder = st.empty()

# if mode == "Review":
#     prompts = {"Summarize": "summarize", "Novelty": "novelty", "Review": "review"}
#     if st.button(f"Run {mode}", key=f"run_{mode}_{doc_id}"):
#         ans, elapsed = bot.ask(
#             prompts[mode], 
#             max_new_tokens=st.session_state.max_new_tokens, 
#             target_words=st.session_state.target_words, 
#             verbosity=st.session_state.verbosity, 
#             strict=st.session_state.strict_grounding,
#             progress_placeholder=progress_placeholder
#         )
#         # store per-doc history
#         doc_history.append({"type": mode, "content": ans, "latency": elapsed})

# # -----------------------------
# # Conversation History
# # -----------------------------
# st.markdown("---")
# st.subheader("📜 Conversation History")
# if not doc_history:
#     st.info("No results generated yet for this document.")
# else:
#     for i, item in enumerate(reversed(doc_history)):
#         lang = item.get("lang", "en")
#         lang_tag = f" [{lang.upper()}]" if lang and lang != "en" else ""
#         base_type = item.get("type", "Result")
#         if "(Translated" in base_type:
#             base_type = base_type.split(" (Translated")[0] + base_type[base_type.find(" (Translated"):]
#         title = f"{base_type}{lang_tag}"
#         if item.get("question"):
#             title += f": {item['question'][:60]}..."

#         with st.expander(title, expanded=(i == 0)):
#             # For Q&A items, content already shows Question/Answer—skip separate "**Q:**" line
#             if item.get("question") and not item.get("type", "").lower().startswith("q&a"):
#                 st.markdown(f"**Q:** {item['question']}")
#             st.markdown(item["content"])
#             if "latency" in item:
#                 st.caption(f"Generated in {item['latency']:.2f} seconds.")

#             if lang == "en":
#                 if st.checkbox("Translate this result?", key=f"translate_cb_{doc_id}_{i}"):
#                     lang_options = {name: code for code, name in Translator.LANG_MAP.items() if code != "en"}
#                     lang_name = st.selectbox("Choose language", list(lang_options.keys()), key=f"lang_select_{doc_id}_{i}")
#                     if st.button("Run Translation", key=f"translate_btn_{doc_id}_{i}"):
#                         target_lang = lang_options[lang_name]
#                         base_type_now = item.get("type", "Result").split(" (Translated")[0]
#                         question_text = item.get("question")
#                         source_lang = item.get("lang", "en") or "en"

#                         if not getattr(translator, "available", False):
#                             how = "Ensure local M2M100 path and install: transformers, sentencepiece, torch"
#                             extra = f" ({getattr(translator, 'last_error', 'unknown error')})"
#                             st.error(f"Translation model not available. {how}{extra}")
#                         else:
#                             try:
#                                 with st.spinner(f"Translating to {lang_name}..."):
#                                     src_text = strip_artifacts(item["content"])
#                                     translated_text = translator.translate(
#                                         src_text, target_lang=target_lang, source_lang=source_lang
#                                     )
#                                     translated_text = strip_artifacts(translated_text)
#                                 if not isinstance(translated_text, str) or translated_text.strip() == item["content"].strip():
#                                     st.error("Translation failed or unchanged output. Check local model path.")
#                                 else:
#                                     st.session_state.history[doc_id].append({
#                                         "type": f"{base_type_now} (Translated — {lang_name.upper()})",
#                                         "question": question_text,
#                                         "content": translated_text,
#                                         "latency": item.get("latency"),
#                                         "lang": target_lang,
#                                     })
#                                     st.success(f"Added {lang_name} translation.")
#                             except Exception as ex:
#                                 st.error(f"Translation error: {ex}")

# # -----------------------------
# # Stop & Export (below Conversation History)
# # -----------------------------
# st.markdown("---")
# col_a, col_b = st.columns([1, 1])
# with col_a:
#     if not st.session_state["export_mode"]:
#         if st.button("🛑 Stop & Export"):
#             st.session_state["export_mode"] = True
#             st.rerun()
# with col_b:
#     if st.session_state["export_mode"]:
#         if st.button("⬅ Back to App"):
#             st.session_state["export_mode"] = False
#             st.rerun()

# if st.session_state["export_mode"]:
#     st.subheader("📦 Export Conversation History")
#     if not doc_history:
#         st.info("No items to export yet.")
#         st.stop()

#     export_fmt = st.selectbox("Choose export format", ["txt", "docx", "pdf"], index=0)
#     if export_fmt == "pdf":
#         try:
#             import reportlab  # type: ignore
#         except Exception:
#             st.error("PDF export requires 'reportlab'. Install with: pip install reportlab")

#     st.markdown("Download files below. Filenames are based on the uploaded document name.")

#     # Individual downloads
#     for idx, item in enumerate(doc_history, start=1):
#         suffix = _export_suffix(item)
#         fname = f"{doc_base}_{suffix}.{export_fmt}"
#         data_bytes = _build_bytes_for_format(item, export_fmt)
#         if data_bytes is None:
#             continue
#         st.download_button(
#             label=f"⬇ Download — {fname}",
#             data=data_bytes,
#             file_name=fname,
#             mime=(
#                 "text/plain" if export_fmt == "txt"
#                 else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#                 if export_fmt == "docx"
#                 else "application/pdf"
#             ),
#             key=f"dl_item_{idx}_{export_fmt}"
#         )

#     # Zip all
#     zip_buf = io.BytesIO()
#     with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
#         for idx, item in enumerate(doc_history, start=1):
#             suffix = _export_suffix(item)
#             fname = f"{doc_base}_{suffix}.{export_fmt}"
#             data_bytes = _build_bytes_for_format(item, export_fmt)
#             if data_bytes is None:
#                 continue
#             zf.writestr(fname, data_bytes)
#     zip_buf.seek(0)
#     st.download_button(
#         label=f"⬇ Download ALL as ZIP ({export_fmt})",
#         data=zip_buf,
#         file_name=f"{doc_base}_history_{export_fmt}.zip",
#         mime="application/zip",
#         key=f"dl_zip_{export_fmt}"
#     )
#     st.stop()

# # -----------------------------
# # Evidence Panel
# # -----------------------------
# if st.session_state.get("show_evidence") and st.session_state.get("__last_evidence__"):
#     st.markdown("---")
#     st.subheader("🔍 Evidence (for the latest result)")
#     for i, e in enumerate(st.session_state["__last_evidence__"], 1):
#         section = e.get("section")
#         page_val = e.get("page")
#         page = f"p.{page_val}" if page_val not in (None, "", "N/A") else None
#         loc_parts = [part for part in (section, page) if part]
#         loc_str = f" ({', '.join(loc_parts)})" if loc_parts else ""
#         with st.expander(f"**[E{i}]**{loc_str} - Chunk ID: {e.get('id', 'N/A')}"):
#             st.caption(e.get("text", ""))





#app.py

import os
import re
import io
import zipfile
import unicodedata
import streamlit as st
from typing import List, Dict, Optional
import sqlalchemy
from sqlalchemy import create_engine, text

# আপনার কনফিগ এবং মডিউল ইমপোর্ট
from config.settings import UPLOAD_DIR, SUPPORTED_EXTENSIONS
from modules.doc_loader import extract_pages
from modules.toc_extractor import extract_toc
from modules.chatbot import Chatbot
from modules.translator import Translator

# -----------------------------
# Database Helper (NEW FEATURE)
# -----------------------------
def init_db():
    """Initializes the database connection and creates the table if not exists."""
    # Docker-compose থেকে DATABASE_URL নিবে, না পেলে লোকালহোস্ট ট্রাই করবে
    db_url = os.getenv("DATABASE_URL", "postgresql://myuser:mypassword@localhost:5432/feedback_db")
    try:
        engine = create_engine(db_url)
        with engine.connect() as conn:
            # টেবিল তৈরি করা (যদি না থাকে)
            conn.execute(text("""
                CREATE TABLE IF NOT EXISTS feedback (
                    id SERIAL PRIMARY KEY,
                    doc_id TEXT,
                    user_query TEXT,
                    rating INT,
                    comments TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """))
            conn.commit()
        return engine
    except Exception as e:
        # ডাটাবেস কানেক্ট না হলে অ্যাপ যেন ক্র্যাশ না করে, তাই এরর প্রিন্ট করে pass করছি
        print(f"⚠️ Database connection warning: {e}")
        return None

# -----------------------------
# Helpers
# -----------------------------
def find_section_for_chunk(chunk_text: str, toc: List[Dict], full_text: str) -> str:
    chunk_pos = full_text.find(chunk_text)
    if chunk_pos == -1:
        return "Unknown"
    current_section = "Introduction"
    for section in toc:
        title = section.get("title", "")
        try:
            pos = full_text.lower().find(title.lower())
            if pos != -1 and pos <= chunk_pos:
                current_section = title
            else:
                break
        except Exception:
            continue
    return current_section

def _safe_build_index(doc_id: str, file_path: str, toc: List[Dict]):
    from modules.vectorstore import VectorStore
    from modules.embedder import Embedder
    store, embedder = VectorStore(doc_id), Embedder()
    pages = extract_pages(file_path)
    full_text = "\n\n".join(p_text for _, p_text in pages)
    chunk_dicts, texts_to_embed = [], []
    for page_num, page_text in pages:
        for chunk in store.split(page_text):
            section_name = find_section_for_chunk(chunk, toc, full_text)
            chunk_dicts.append({"text": chunk, "page": page_num, "section": section_name})
            texts_to_embed.append(chunk)
    if not texts_to_embed:
        raise ValueError("Document is empty or could not be chunked.")
    vectors = embedder.encode(texts_to_embed)
    store.build_index(vectors, chunk_dicts)
    return len(chunk_dicts)

@st.cache_resource
def get_chatbot(doc_id, file_path):
    return Chatbot(doc_id, file_path)

@st.cache_resource
def get_translator():
    model_path = os.getenv("TRANSLATION_MODEL", "models/m2m100_418M")
    return Translator(model_path)

def render_toc(toc, level=0):
    for entry in toc:
        indent = "&nbsp;" * 4 * (entry.get("level", 0))
        st.sidebar.markdown(
            f"{indent}- **{entry['title']}**" + (f" (p.{entry.get('page')})" if entry.get("page") else ""),
            unsafe_allow_html=True,
        )
        if entry.get("children"):
            render_toc(entry["children"], level + 1)

# --- Output cleaning / sanitizing ---
_DROP_PATTERNS = [
    r"\bThe provided evidence does not contain\b.*",
    r"\bThe provided evidence does not\b.*",
]

def clean_answer(raw: str) -> str:
    if not isinstance(raw, str) or not raw.strip():
        return raw
    txt = raw.replace("[/ ", "").replace("[/]", "").replace("[ /]", "")
    out_lines = []
    for ln in txt.splitlines():
        if ln.lstrip().startswith("|"):
            out_lines.append(ln)
            continue
        if any(re.search(p, ln, re.I) for p in _DROP_PATTERNS):
            continue
        ln = re.sub(r"^\s*\*\*(.*?)\s*\*\*\s*", r"\1", ln)
        out_lines.append(ln)
    txt = "\n".join(out_lines)
    txt = re.sub(r"\n{3,}", "\n\n", txt).strip()
    return txt

_CLEAN_TAGS = [
    r"\[E\d+\]",
    r"\[/\s*",
]

def strip_artifacts(text: str) -> str:
    if not isinstance(text, str):
        return text
    out = text
    for p in _CLEAN_TAGS:
        out = re.sub(p, "", out)
    out = re.sub(r"\s{2,}", " ", out)
    out = re.sub(r"\n{3,}", "\n\n", out)
    return out.strip()

def _slugify(text: str, maxlen: int = 80) -> str:
    if not text:
        return ""
    value = unicodedata.normalize("NFKD", text).encode("ascii", "ignore").decode("ascii")
    value = re.sub(r"[^\w\s-]+", "", value).strip().lower()
    value = re.sub(r"[-\s]+", "_", value)
    return value[:maxlen] if maxlen else value

def _export_suffix(item: Dict) -> str:
    typ = (item.get("type") or "").lower()
    lang = (item.get("lang") or "en").lower()

    if typ.startswith("q&a"):
        base = "qa"
    elif typ.startswith("summarize") or "summary" in typ:
        base = "summary"
    elif typ.startswith("novelty"):
        base = "novelty"
    elif typ.startswith("review"):
        base = "review"
    else:
        base = _slugify(item.get("type") or "result")

    if "translated" in typ:
        m = re.search(r"translated\s*—\s*([A-Z]+)", item.get("type", ""), re.I)
        if m:
            lang_name = m.group(1).lower()
        else:
            lang_name = lang
        return f"{base}_{lang_name}_translated"

    return base

def _item_to_text(item: Dict) -> str:
    parts = []
    if item.get("type", "").lower().startswith("q&a"):
        parts.append(item.get("content", ""))
    else:
        if item.get("question"):
            parts.append(f"Q: {item['question']}")
        parts.append(item.get("content", ""))
    if "latency" in item:
        parts.append(f"\n---\nGenerated in {item['latency']:.2f} seconds.")
    return "\n\n".join([p for p in parts if p])

def _build_txt_bytes(item: Dict) -> bytes:
    return _item_to_text(item).encode("utf-8")

def _build_docx_bytes(item: Dict) -> bytes:
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    if item.get("type"):
        p = doc.add_paragraph()
        run = p.add_run(item["type"])
        run.bold = True
        run.font.size = Pt(14)
    if item.get("question") and not item.get("type", "").lower().startswith("q&a"):
        doc.add_paragraph(f"Q: {item['question']}")
    for line in _item_to_text(item).splitlines():
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

def _build_pdf_bytes(item: Dict) -> Optional[bytes]:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import cm
    except Exception:
        return None
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4
    x, y = 2 * cm, height - 2 * cm

    title = item.get("type", "Result")
    c.setFont("Helvetica-Bold", 12)
    c.drawString(x, y, title)
    y -= 0.8 * cm

    if item.get("question") and not item.get("type", "").lower().startswith("q&a"):
        c.setFont("Helvetica-Bold", 10)
        c.drawString(x, y, f"Q: {item['question']}")
        y -= 0.7 * cm

    c.setFont("Helvetica", 10)
    textobj = c.beginText(x, y)
    for line in _item_to_text(item).splitlines():
        textobj.textLine(line if line else " ")
        if textobj.getY() < 2 * cm:
            c.drawText(textobj)
            c.showPage()
            textobj = c.beginText(x, height - 2 * cm)
            c.setFont("Helvetica", 10)
    c.drawText(textobj)
    c.showPage()
    c.save()
    buf.seek(0)
    return buf.read()

def _build_bytes_for_format(item: Dict, fmt: str) -> Optional[bytes]:
    fmt = fmt.lower()
    if fmt == "txt":
        return _build_txt_bytes(item)
    elif fmt == "docx":
        return _build_docx_bytes(item)
    elif fmt == "pdf":
        return _build_pdf_bytes(item)
    else:
        return _build_txt_bytes(item)

def format_qa_block(question: str, answer: str) -> str:
    q = (question or "").strip()
    a = (answer or "").strip()
    a = re.sub(r"^\s*(Q|Question)\s*:\s*.*\n+", "", a, flags=re.IGNORECASE)
    a = re.sub(r"^\s*(A|Answer)\s*:\s*", "", a, flags=re.IGNORECASE)
    escaped_q = re.escape(q.rstrip("?.!\"'`*:_- \t"))
    leading_q_patterns = [
        rf"^\s*\*\*{escaped_q}\*\*\s*[:\-–—]?\s*",
        rf'^\s*"{escaped_q}"\s*[:\-–—]?\s*',
        rf"^\s*{escaped_q}\s*[?.!:]*\s*[:\-–—]?\s*",
    ]
    for pat in leading_q_patterns:
        new_a = re.sub(pat, "", a, flags=re.IGNORECASE)
        if new_a != a:
            a = new_a
            break
    def _norm(s: str) -> str:
        s = s.lower().strip()
        s = re.sub(r"[^\w\s]", " ", s)
        s = re.sub(r"\s+", " ", s)
        return s
    first_sentence_match = re.split(r"(?<=[.!?])\s+", a, maxsplit=1)
    first_sent = first_sentence_match[0] if first_sentence_match else a
    if _norm(first_sent) == _norm(q) or _norm(first_sent.rstrip("?.!:")) == _norm(q):
        a = a[len(first_sent):].lstrip()
    a = re.sub(r"\n{3,}", "\n\n", a).strip()
    return f"**Question**\n\n{q}\n\n**Answer**\n\n{a}\n"

# -----------------------------
# Page & State
# -----------------------------
st.set_page_config(page_title="Smart Document Analysis", layout="wide")
st.title("📄 Smart Document Analyst")

st.session_state.setdefault("docs", [])
st.session_state.setdefault("history", {})
st.session_state.setdefault("active_doc_id", None)
st.session_state.setdefault("export_mode", False)
st.session_state.setdefault("active_doc_base", None)

translator = get_translator()

# -----------------------------
# Sidebar Controls
# -----------------------------
st.sidebar.subheader("⚙️ Generation Controls")
st.session_state.setdefault("max_new_tokens", 900)
st.session_state.setdefault("target_words", 600)
st.session_state.setdefault("verbosity", "detailed")
st.session_state.setdefault("strict_grounding", True)
st.session_state.setdefault("show_evidence", True)

st.session_state["max_new_tokens"] = st.sidebar.number_input(
    "Max new tokens", min_value=128, max_value=4096, step=64, value=st.session_state["max_new_tokens"]
)
st.session_state["target_words"] = st.sidebar.slider(
    "Target length (words)", min_value=150, max_value=1500, step=50, value=st.session_state["target_words"]
)
st.session_state["verbosity"] = st.sidebar.selectbox(
    "Writing style", ["concise", "balanced", "detailed", "exhaustive"],
    index=["concise", "balanced", "detailed", "exhaustive"].index(st.session_state["verbosity"]),
)
st.session_state["strict_grounding"] = st.sidebar.checkbox(
    "Strict grounding (evidence-only, no invention)", value=st.session_state["strict_grounding"]
)
st.session_state["show_evidence"] = st.sidebar.checkbox("Show evidence panel", value=st.session_state["show_evidence"])

# -----------------------------
# Uploader & Indexing
# -----------------------------
uploaded = st.file_uploader("Upload your document", type=[e.lstrip(".") for e in SUPPORTED_EXTENSIONS])
if uploaded:
    doc_id, _ext = os.path.splitext(uploaded.name)
    st.session_state["active_doc_base"] = doc_id
    if not any(d["doc_id"] == doc_id for d in st.session_state.docs):
        path = os.path.join(UPLOAD_DIR, uploaded.name)
        os.makedirs(UPLOAD_DIR, exist_ok=True)
        with open(path, "wb") as f:
            f.write(uploaded.getbuffer())
        with st.spinner(f"Processing and indexing {uploaded.name}..."):
            toc = extract_toc(path)
            st.session_state.docs.append({"file_path": path, "doc_id": doc_id, "toc": toc})
            st.session_state.history[doc_id] = []
            try:
                _safe_build_index(doc_id, path, toc)
                st.sidebar.success(f"✅ Indexed {uploaded.name}.")
                st.session_state.active_doc_id = doc_id
            except Exception as e:
                st.sidebar.error(f"Index build failed: {e}")
    else:
        if st.session_state.active_doc_id != doc_id:
            st.session_state.active_doc_id = doc_id

# -----------------------------
# Main App
# -----------------------------
if not st.session_state.active_doc_id:
    st.info("Please upload a document to begin.")
    st.stop()

doc_id = st.session_state.active_doc_id
try:
    current_doc = next(d for d in st.session_state.docs if d["doc_id"] == doc_id)
except StopIteration:
    st.error("Active document not found. Please re-upload.", icon="🚨")
    st.stop()

if not st.session_state.get("active_doc_base"):
    st.session_state["active_doc_base"] = current_doc["doc_id"]

file_path = current_doc["file_path"]
doc_history = st.session_state.history.setdefault(doc_id, [])
doc_base = st.session_state["active_doc_base"]

# -----------------------------
# Sidebar: Table of Contents (ONLY here)
# -----------------------------
if st.sidebar.checkbox("📚 Show Table of Contents", value=True, key="show_toc"):
    toc = current_doc.get("toc", [])
    if toc:
        st.sidebar.subheader("📚 Table of Contents")
        render_toc(toc)
    else:
        st.sidebar.info("No Table of Contents extracted for this document.")

# -----------------------------
# Generate
# -----------------------------
st.subheader("💬 Chat & Generate")
mode = st.radio("Select Mode", ["General Q&A", "Summarize", "Novelty", "Review"], horizontal=True, key=f"mode_{doc_id}")
bot = get_chatbot(doc_id, file_path)

progress_placeholder = st.empty()

# --- Logic for Q&A ---
if mode == "General Q&A":
    q = st.text_input("Your question:", key=f"q_{doc_id}")
    if st.button("Ask", key=f"ask_{doc_id}") and q:
        with st.spinner("🤖 Thinking..."):
            ans, elapsed = bot.ask(
                q,
                max_new_tokens=st.session_state.max_new_tokens,
                target_words=st.session_state.target_words,
                verbosity=st.session_state.verbosity,
                strict=st.session_state.strict_grounding,
                progress_placeholder=progress_placeholder
            )
        doc_history.append({"type": "Q&A", "question": q, "content": ans, "latency": elapsed})

# --- Logic for Summarize/Novelty/Review ---
elif mode in ["Summarize", "Novelty", "Review"]:
    prompts = {"Summarize": "summarize", "Novelty": "novelty", "Review": "review"}
    if st.button(f"Run {mode}", key=f"run_{mode}_{doc_id}"):
        ans, elapsed = bot.ask(
            prompts[mode],
            max_new_tokens=st.session_state.max_new_tokens,
            target_words=st.session_state.target_words,
            verbosity=st.session_state.verbosity,
            strict=st.session_state.strict_grounding,
            progress_placeholder=progress_placeholder
        )
        doc_history.append({"type": mode, "content": ans, "latency": elapsed})

# -----------------------------
# Conversation History
# -----------------------------
st.markdown("---")
st.subheader("📜 Conversation History")
if not doc_history:
    st.info("No results generated yet for this document.")
else:
    for i, item in enumerate(reversed(doc_history)):
        lang = item.get("lang", "en")
        lang_tag = f" [{lang.upper()}]" if lang and lang != "en" else ""
        base_type = item.get("type", "Result")
        if "(Translated" in base_type:
            base_type = base_type.split(" (Translated")[0] + base_type[base_type.find(" (Translated"):]
        title = f"{base_type}{lang_tag}"
        if item.get("question"):
            title += f": {item['question'][:60]}..."

        with st.expander(title, expanded=(i == 0)):
            if item.get("question") and not item.get("type", "").lower().startswith("q&a"):
                st.markdown(f"**Q:** {item['question']}")
            st.markdown(item["content"])
            if "latency" in item:
                st.caption(f"Generated in {item['latency']:.2f} seconds.")

            if lang == "en":
                if st.checkbox("Translate this result?", key=f"translate_cb_{doc_id}_{i}"):
                    lang_options = {name: code for code, name in Translator.LANG_MAP.items() if code != "en"}
                    lang_name = st.selectbox("Choose language", list(lang_options.keys()), key=f"lang_select_{doc_id}_{i}")
                    if st.button("Run Translation", key=f"translate_btn_{doc_id}_{i}"):
                        target_lang = lang_options[lang_name]
                        base_type_now = item.get("type", "Result").split(" (Translated")[0]
                        question_text = item.get("question")
                        source_lang = item.get("lang", "en") or "en"

                        if not getattr(translator, "available", False):
                            st.error(f"Translation model not available. ({getattr(translator, 'last_error', 'unknown error')})")
                        else:
                            try:
                                with st.spinner(f"Translating to {lang_name}..."):
                                    src_text = strip_artifacts(item["content"])
                                    translated_text = translator.translate(
                                        src_text, target_lang=target_lang, source_lang=source_lang
                                    )
                                    translated_text = strip_artifacts(translated_text)
                                if not isinstance(translated_text, str) or translated_text.strip() == item["content"].strip():
                                    st.error("Translation failed or unchanged output.")
                                else:
                                    st.session_state.history[doc_id].append({
                                        "type": f"{base_type_now} (Translated — {lang_name.upper()})",
                                        "question": question_text,
                                        "content": translated_text,
                                        "latency": item.get("latency"),
                                        "lang": target_lang,
                                    })
                                    st.success(f"Added {lang_name} translation.")
                                    st.rerun()
                            except Exception as ex:
                                st.error(f"Translation error: {ex}")

# -----------------------------
# Stop & Export
# -----------------------------
st.markdown("---")
col_a, col_b = st.columns([1, 1])
with col_a:
    if not st.session_state["export_mode"]:
        if st.button("🛑 Stop & Export"):
            st.session_state["export_mode"] = True
            st.rerun()
with col_b:
    if st.session_state["export_mode"]:
        if st.button("⬅ Back to App"):
            st.session_state["export_mode"] = False
            st.rerun()

if st.session_state["export_mode"]:
    st.subheader("📦 Export Conversation History")
    if not doc_history:
        st.info("No items to export yet.")
    else:
        export_fmt = st.selectbox("Choose export format", ["txt", "docx", "pdf"], index=0)
        st.markdown("Download files below. Filenames are based on the uploaded document name.")

        # Zip all
        zip_buf = io.BytesIO()
        with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
            for idx, item in enumerate(doc_history, start=1):
                suffix = _export_suffix(item)
                fname = f"{doc_base}_{suffix}.{export_fmt}"
                data_bytes = _build_bytes_for_format(item, export_fmt)
                if data_bytes:
                    zf.writestr(fname, data_bytes)
        zip_buf.seek(0)
        st.download_button(
            label=f"⬇ Download ALL as ZIP ({export_fmt})",
            data=zip_buf,
            file_name=f"{doc_base}_history_{export_fmt}.zip",
            mime="application/zip",
            key=f"dl_zip_{export_fmt}"
        )
    st.stop()

# -----------------------------
# NEW FEATURE: Feedback System (Database)
# -----------------------------
st.markdown("---")
st.subheader("⭐ Give Feedback (Saved to DB)")
with st.form("feedback_form"):
    col_f1, col_f2 = st.columns([1, 2])
    with col_f1:
        f_rating = st.slider("Rate the output (1-5)", 1, 5, 5)
    with col_f2:
        f_comment = st.text_input("Comment (Optional)", placeholder="Was the answer accurate?")
    
    submitted = st.form_submit_button("Submit Feedback")
    if submitted:
        engine = init_db()
        if engine:
            try:
                last_query = "N/A"
                if doc_history:
                    last_item = doc_history[-1]
                    last_query = last_item.get("question") or last_item.get("type") or "Unknown"

                with engine.connect() as conn:
                    query = text("INSERT INTO feedback (doc_id, user_query, rating, comments) VALUES (:d, :q, :r, :c)")
                    conn.execute(query, {"d": doc_id, "q": last_query, "r": f_rating, "c": f_comment})
                    conn.commit()
                st.success("✅ Feedback saved to PostgreSQL Database!")
            except Exception as e:
                st.error(f"Failed to save feedback: {e}")
        else:
            st.warning("Database not connected. Is the 'db' container running?")

# -----------------------------
# Evidence Panel
# -----------------------------
if st.session_state.get("show_evidence") and st.session_state.get("__last_evidence__"):
    st.markdown("---")
    st.subheader("🔍 Evidence (for the latest result)")
    for i, e in enumerate(st.session_state["__last_evidence__"], 1):
        section = e.get("section")
        page_val = e.get("page")
        page = f"p.{page_val}" if page_val not in (None, "", "N/A") else None
        loc_parts = [part for part in (section, page) if part]
        loc_str = f" ({', '.join(loc_parts)})" if loc_parts else ""
        with st.expander(f"**[E{i}]**{loc_str} - Chunk ID: {e.get('id', 'N/A')}"):
            st.caption(e.get("text", ""))